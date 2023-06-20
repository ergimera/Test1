import streamlit as st
from docx import Document

def create_docx():
    # Create a new Document object
    doc = Document()

    # Add content to the document
    doc.add_heading('My Document', level=1)
    doc.add_paragraph('This is some sample content.')

    # Save the document
    doc.save('output.docx')
